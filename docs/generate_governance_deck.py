"""
Generate a Word document presenting the Azure Operations Agent
as a Governance Agent for the STU M1 team meeting / sales team manager.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Azure_Governance_Agent_Use_Cases.docx")

# Try alternate path if file is locked
FALLBACK_PATH = os.path.join(os.path.dirname(__file__), "Azure_Governance_Agent_Use_Cases_v2.docx")

# --- Helpers ----------------------------------------------------------------

BRAND_BLUE = RGBColor(0x00, 0x78, 0xD4)   # Microsoft blue
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex: str):
    """Set the background colour of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "0078D4")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F0F0")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# --- Document creation -------------------------------------------------------

doc = Document()

# Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Segoe UI"
font.size = Pt(10.5)
font.color.rgb = DARK_GRAY

# Reduce spacing after paragraphs
style.paragraph_format.space_after = Pt(6)

# ============================================================================
# TITLE PAGE
# ============================================================================
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Azure Governance Agent")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = BRAND_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Solution Accelerator for AI-Powered Azure Governance")
run.font.size = Pt(16)
run.font.color.rgb = DARK_GRAY

doc.add_paragraph("")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Use-Case Presentation for STU M1 Team Meeting\nSales Team Manager Review")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph("")
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("March 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================================
# EXECUTIVE SUMMARY
# ============================================================================
h = doc.add_heading("Executive Summary", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

doc.add_paragraph(
    "The Azure Governance Agent is a solution accelerator that provides a ready-to-deploy, "
    "extensible framework for building AI-powered Azure governance agents. Built with the "
    "Microsoft Agent Framework and Model Context Protocol (MCP), it enables organizations to "
    "monitor, manage, and enforce governance across their Azure environment using natural language — "
    "all through a conversational chat interface with rich visual dashboards."
)
doc.add_paragraph(
    "This document outlines the key governance use cases, the value proposition for customers, "
    "and how the solution accelerator empowers Azure admins to build custom governance agents "
    "tailored to their organization's needs."
)

doc.add_paragraph("")

# ============================================================================
# WHAT IS A SOLUTION ACCELERATOR?
# ============================================================================
h = doc.add_heading("What is a Solution Accelerator?", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

doc.add_paragraph(
    "A solution accelerator is a pre-built, production-ready codebase that customers and partners "
    "can use as a starting point to rapidly build and customize their own solution. Unlike a finished "
    "product, a solution accelerator is designed to be forked, extended, and tailored to specific "
    "organizational requirements."
)

doc.add_heading("What Ships Today", level=2)
doc.add_paragraph(
    "The accelerator ships with two fully functional governance agents out of the box:"
)
built_in = [
    ("Resource Discovery & Unused Resource Agent",
     "Discovers all Azure resources across subscriptions using Resource Graph (KQL), "
     "identifies orphaned disks, unattached public IPs, idle VMs, unused NICs and NSGs, "
     "and generates interactive reports with cost-saving recommendations."),
    ("Policy Agent",
     "Lists policy assignments, checks compliance status, and authors custom Azure Policy "
     "definitions — including ready-to-deploy deny-public-IP and restrict-locations policies "
     "with CLI commands for immediate deployment."),
]
for title_text, desc in built_in:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("")

doc.add_heading("Extensibility: Build Your Own Agents", level=2)
doc.add_paragraph(
    "The real power of the accelerator is its extensible architecture. Azure admins and "
    "engineering teams can add additional governance agents using the same framework — "
    "without modifying the core platform. The MCP tool-based architecture makes it simple "
    "to plug in new capabilities."
)

extension_examples = [
    ("Custom Visual Dashboards",
     "Build organization-specific dashboards and reports with custom charts, branding, and KPIs."),
    ("Email & Notification Integration",
     "Connect to Exchange Online, SendGrid, Azure Logic Apps, or Microsoft Graph to send "
     "governance alerts, reports, and action summaries to stakeholders."),
    ("ServiceNow Integration",
     "Automatically create ServiceNow incidents or change requests when governance violations "
     "are detected — orphaned resources, policy non-compliance, cost threshold breaches."),
    ("Tagging & Compliance Agents",
     "Enforce tagging standards, detect untagged resources, and auto-remediate tagging gaps."),
    ("Cost Threshold Alerting",
     "Build agents that monitor spend against budgets and trigger escalation workflows "
     "when thresholds are exceeded."),
    ("Security Posture Agent",
     "Extend with Microsoft Defender for Cloud integration to surface security recommendations "
     "and track Secure Score improvements."),
]
for title_text, desc in extension_examples:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("")

doc.add_heading("How Extension Works", level=2)
extension_steps = [
    "Create a new MCP tool file (e.g., servicenow_tools.py) with your custom integrations.",
    "Register the tools with the MCP server — they are automatically available to the agent.",
    "Optionally add a new specialist agent with domain-specific instructions in the orchestrator.",
    "The HandoffBuilder automatically routes relevant questions to your new agent.",
    "Custom visuals can be added as new report templates in the report tools.",
]
for step in extension_steps:
    doc.add_paragraph(step, style="List Number")

doc.add_paragraph("")

# ============================================================================
# THE PROBLEM
# ============================================================================
h = doc.add_heading("The Problem: Cloud Governance at Scale", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

problems = [
    ("Visibility Gap", "Organizations with hundreds of subscriptions struggle to get a unified view of resources, costs, and compliance status."),
    ("Cost Overruns", "Orphaned disks, idle VMs, and unattached public IPs silently rack up charges — often going unnoticed for months."),
    ("Policy Drift", "Policies are defined but not actively monitored. Non-compliant resources proliferate without timely remediation."),
    ("Manual Toil", "Cloud ops teams spend hours in the Azure Portal clicking through blades, running CLI scripts, and assembling reports manually."),
    ("Slow Incident Response", "Identifying unhealthy resources, reviewing activity logs, and correlating metrics requires deep Azure expertise and multiple tools."),
]
for title_text, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    run.font.size = Pt(10.5)
    p.add_run(desc)

doc.add_paragraph("")

# ============================================================================
# SOLUTION OVERVIEW
# ============================================================================
h = doc.add_heading("Solution Accelerator: Azure Governance Agent", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

doc.add_paragraph(
    "The Azure Governance Agent solution accelerator replaces manual portal workflows with a single, "
    "conversational interface that understands natural language queries and executes Azure operations "
    "on behalf of the user — securely, using their own Entra ID identity and RBAC permissions. "
    "It ships with built-in agents for resource discovery, unused resource management, and policy "
    "authoring, and is designed to be extended with additional agents for email notifications, "
    "ServiceNow ticketing, custom dashboards, and any other governance workflow."
)

doc.add_heading("Architecture Highlights", level=2)
arch_rows = [
    ("React SPA", "Modern chat & dashboard UI with Entra ID authentication (MSAL)"),
    ("FastAPI Backend", "JWT-validated API layer with AI agent orchestration and NDJSON streaming"),
    ("MCP Server", "Tool server with 25+ Azure integrations (Resource Graph, Monitor, Cost Management, Policy, ARM)"),
    ("Microsoft Agent Framework", "Multi-agent orchestration with automatic handoff between Ops and Policy specialists"),
    ("Token Passthrough", "User's Entra ID token flows securely from browser → API → MCP → Azure APIs (no stored credentials)"),
]
add_styled_table(doc, ["Component", "Description"], arch_rows, col_widths=[5, 12])

doc.add_paragraph("")

# ============================================================================
# USE CASES
# ============================================================================
h = doc.add_heading("Governance Use Cases", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

# --- UC 1 ---
doc.add_heading("Use Case 1: Resource Visibility & Discovery", level=2)
doc.add_paragraph(
    "Ask the agent in plain English to discover all resources across subscriptions. "
    "The agent uses Azure Resource Graph (KQL) to query resources in real time."
)
uc1_examples = [
    '"Show me all VMs in East US that are deallocated"',
    '"List all storage accounts across my subscriptions"',
    '"How many resources do I have by type and region?"',
    '"Find all resources tagged with CostCenter=Marketing"',
]
for ex in uc1_examples:
    doc.add_paragraph(ex, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Business Value: ")
run.bold = True
p.add_run(
    "Instant, cross-subscription inventory without portal navigation. "
    "Saves 30-60 minutes per ad-hoc resource lookup."
)

doc.add_paragraph("")

# --- UC 2 ---
doc.add_heading("Use Case 2: Cost Governance & Optimization", level=2)
doc.add_paragraph(
    "The agent provides real-time cost breakdowns, identifies top spend areas, "
    "and surfaces Azure Advisor cost recommendations — all through chat."
)
uc2_examples = [
    '"What is my total Azure spend this month?"',
    '"Break down costs by resource group for the last 30 days"',
    '"Show me my top 10 most expensive resources"',
    '"What does Azure Advisor recommend for cost savings?"',
    '"Find orphaned disks and unused public IPs that are costing money"',
    '"Generate a cost report and email it to the finance team"',
]
for ex in uc2_examples:
    doc.add_paragraph(ex, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Business Value: ")
run.bold = True
p.add_run(
    "Proactive cost governance. Customers typically discover 10-25% in wasted spend "
    "from orphaned and idle resources within the first scan."
)

doc.add_paragraph("")

# --- UC 3 ---
doc.add_heading("Use Case 3: Policy Compliance & Enforcement", level=2)
doc.add_paragraph(
    "A dedicated Policy Agent handles all Azure Policy operations — listing assignments, "
    "checking compliance, and authoring custom policies with ready-to-deploy CLI commands."
)
uc3_examples = [
    '"List all policy assignments on my subscription"',
    '"What is my policy compliance status?"',
    '"Create a policy to deny public IP addresses"',
    '"Generate a policy that restricts resources to East US and West US only"',
    '"Show me all non-compliant resources for the tagging policy"',
]
for ex in uc3_examples:
    doc.add_paragraph(ex, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Business Value: ")
run.bold = True
p.add_run(
    "Continuous compliance monitoring and policy creation in minutes instead of hours. "
    "No need to memorize policy JSON schema or ARM API paths."
)

doc.add_paragraph("")

# --- UC 4 ---
doc.add_heading("Use Case 4: Operational Health & Monitoring", level=2)
doc.add_paragraph(
    "Query Azure Monitor metrics, check resource health, review activity logs, "
    "and detect idle resources — all conversationally."
)
uc4_examples = [
    '"Is my VM healthy? What are its CPU and memory metrics?"',
    '"Show me activity logs for the last 24 hours"',
    '"Are there any idle resources wasting money?"',
    '"Check the health of all resources in the Production resource group"',
    '"What metric alerts are configured on my subscription?"',
]
for ex in uc4_examples:
    doc.add_paragraph(ex, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Business Value: ")
run.bold = True
p.add_run(
    "Faster incident triage and proactive health monitoring. "
    "Reduces Mean-Time-To-Detect (MTTD) by eliminating portal context-switching."
)

doc.add_paragraph("")

# --- UC 5 ---
doc.add_heading("Use Case 5: Automated Reporting & Dashboards", level=2)
doc.add_paragraph(
    "The agent generates interactive HTML reports and dashboards on demand — "
    "resource inventories, cost visualizations, and comprehensive environment overviews."
)
uc5_examples = [
    '"Generate a dashboard report for my Azure environment"',
    '"Create a resource report for all VMs"',
    '"Build a cost report for the last quarter"',
    '"Generate a report of orphaned resources and email it to the ops team"',
]
for ex in uc5_examples:
    doc.add_paragraph(ex, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Business Value: ")
run.bold = True
p.add_run(
    "Eliminates manual report assembly. Stakeholders (finance, security, ops) "
    "get self-service, data-rich reports in seconds."
)

doc.add_paragraph("")

# --- UC 6 ---
doc.add_heading("Use Case 6: Resource Management & VM Operations", level=2)
doc.add_paragraph(
    "Perform day-to-day operational tasks through chat — start/stop VMs, "
    "manage tags, and list subscription details."
)
uc6_examples = [
    '"Stop all VMs in the Dev resource group"',
    '"Tag all resources in Production with Environment=Prod"',
    '"List all my Azure subscriptions"',
    '"Restart the web server VM"',
]
for ex in uc6_examples:
    doc.add_paragraph(ex, style="List Bullet")

p = doc.add_paragraph()
run = p.add_run("Business Value: ")
run.bold = True
p.add_run(
    "Operational efficiency — common tasks that take 5-10 clicks in the portal "
    "are reduced to a single chat message."
)

doc.add_page_break()

# ============================================================================
# DIFFERENTIATORS
# ============================================================================
h = doc.add_heading("Key Differentiators", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

diff_rows = [
    ("Solution Accelerator Model", "Not a black-box product — a fully open, forkable codebase that customers and partners customize to their governance needs."),
    ("Natural Language Interface", "No CLI expertise or portal navigation required. Users ask questions in plain English."),
    ("Zero Credential Storage", "Uses Entra ID token passthrough — the agent never stores credentials. All actions use the user's own RBAC permissions."),
    ("Multi-Agent Architecture", "Specialized agents (Ops, Policy) with intelligent routing. Admins add new agents without modifying the core."),
    ("Pluggable Integrations", "Add email (Exchange/SendGrid), ServiceNow ticketing, Slack/Teams notifications, or any REST API as new MCP tools."),
    ("Real-Time Azure Data", "Every response is backed by live Azure API calls — no stale data or periodic sync."),
    ("Custom Visual Dashboards", "Rich HTML dashboards with charts, tables, and filters — generated on demand, fully customizable per organization."),
    ("Built on Microsoft Stack", "Microsoft Agent Framework, Entra ID, Azure OpenAI, MCP — fully aligned with the Microsoft ecosystem."),
    ("Extensible Tool Architecture", "MCP-based tools are modular. New Azure integrations can be added without changing the agent core."),
]
add_styled_table(doc, ["Differentiator", "Description"], diff_rows, col_widths=[5.5, 11.5])

doc.add_paragraph("")

# ============================================================================
# CUSTOMER VALUE PROPOSITION
# ============================================================================
h = doc.add_heading("Customer Value Proposition", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

value_rows = [
    ("Reduce cloud waste", "Identify orphaned/idle resources and recover 10-25% of wasted spend"),
    ("Accelerate governance", "Policy creation, compliance checks, and remediation in minutes"),
    ("Empower non-experts", "Business users and junior admins can query Azure without CLI/portal expertise"),
    ("Cut reporting time", "Automated, interactive reports replace hours of manual data gathering"),
    ("Improve security posture", "Continuous compliance monitoring with proactive alerts and recommendations"),
    ("Drive Azure adoption", "Lower the learning curve for Azure management, increasing platform stickiness"),
    ("Faster time-to-value", "Solution accelerator gets customers from zero to working governance agent in days, not months"),
    ("Customizable to org needs", "Admins add ServiceNow, email, Teams, or any custom integration without rebuilding the platform"),
]
add_styled_table(doc, ["Value", "Impact"], value_rows, col_widths=[5, 12])

doc.add_paragraph("")

# ============================================================================
# TARGET CUSTOMERS
# ============================================================================
h = doc.add_heading("Target Customer Profiles", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

doc.add_paragraph(
    "This solution is ideal for organizations that need to improve their Azure governance posture:"
)
targets = [
    ("Enterprise Cloud Teams", "Managing 50+ subscriptions with thousands of resources across multiple regions and business units."),
    ("Managed Service Providers (MSPs)", "Overseeing multiple customer tenants and needing efficient cross-tenant governance."),
    ("Finance & FinOps Teams", "Requiring real-time cost visibility, budget tracking, and spend optimization insights."),
    ("Security & Compliance Teams", "Enforcing organizational policies, auditing compliance, and ensuring security baselines."),
    ("DevOps / Platform Engineering", "Automating resource lifecycle management, tagging standards, and operational hygiene."),
]
for title_text, desc in targets:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ============================================================================
# DEMO SCENARIO
# ============================================================================
h = doc.add_heading("Suggested Demo Scenario", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

doc.add_paragraph(
    "A step-by-step demo flow to showcase the governance capabilities in a customer meeting:"
)

demo_steps = [
    ("Login & Auth", 'Open the SPA, sign in with Entra ID. Highlight "zero credential storage" — the agent uses the user\'s own identity.'),
    ("Resource Discovery", '"Show me a summary of all my Azure resources." → Agent returns resource counts by type, region, and resource group.'),
    ("Cost Deep-Dive", '"What is my spend this month by resource group?" → Agent shows cost breakdown. Follow up: "Show my top 5 most expensive resources."'),
    ("Find Waste", '"Find orphaned disks and unused public IPs." → Agent identifies waste. "Generate a report and email it to the ops team."'),
    ("Policy Compliance", '"What is my policy compliance status?" → Agent shows compliant vs non-compliant counts. "Create a policy to deny public IPs."'),
    ("Dashboard Report", '"Generate a full dashboard report for my environment." → Interactive HTML dashboard with charts, tables, and filters.'),
    ("Operational Action", '"Stop all deallocated VMs in the Dev resource group." → Agent confirms before executing, demonstrating safe operational control.'),
]

for i, (step_title, step_desc) in enumerate(demo_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Step {i} — {step_title}: ")
    run.bold = True
    p.add_run(step_desc)

doc.add_paragraph("")

# ============================================================================
# ROADMAP
# ============================================================================
h = doc.add_heading("Roadmap Highlights", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

roadmap_rows = [
    ("Actionable Policy Agent", "Create, assign, and remediate Azure Policies directly from chat", "Planned"),
    ("Automated Cleanup Actions", "Delete orphaned resources, resize idle VMs, stop unused services", "Planned"),
    ("Email Integration", "Real email delivery via Logic Apps, Graph API, or Action Groups", "Planned"),
    ("Enhanced UI", "JSON viewer, policy badges, sortable tables, copy-to-clipboard", "Planned"),
    ("Multi-Tenant Support", "Cross-tenant governance for MSP scenarios", "Under Evaluation"),
]
add_styled_table(doc, ["Feature", "Description", "Status"], roadmap_rows, col_widths=[5, 9, 3])

doc.add_paragraph("")

# ============================================================================
# NEXT STEPS
# ============================================================================
h = doc.add_heading("Next Steps", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

next_steps = [
    "Schedule a live demo with the STU M1 team to walk through the governance scenarios.",
    "Identify 2-3 pilot customers with immediate governance pain points (cost waste, policy drift).",
    "Align with the Azure sales motion — position alongside Azure Advisor, Cost Management, and Azure Policy.",
    "Gather feedback from the sales team on additional use cases and customer objections.",
    "Plan a joint session with engineering to discuss customization for specific customer environments.",
]
for step in next_steps:
    doc.add_paragraph(step, style="List Number")

doc.add_paragraph("")

# ============================================================================
# CONTACT / CLOSING
# ============================================================================
h = doc.add_heading("Contact & Resources", level=1)
for r in h.runs:
    r.font.color.rgb = BRAND_BLUE

doc.add_paragraph("For questions, demo requests, or feedback:")
doc.add_paragraph("Repository: AzureAgent (internal)", style="List Bullet")
doc.add_paragraph("Architecture: React SPA + FastAPI + MCP Server + Microsoft Agent Framework", style="List Bullet")
doc.add_paragraph("Authentication: Microsoft Entra ID (MSAL + token passthrough)", style="List Bullet")

# ============================================================================
# SAVE
# ============================================================================
try:
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
except PermissionError:
    doc.save(FALLBACK_PATH)
    print(f"Original file locked. Document saved to: {FALLBACK_PATH}")
